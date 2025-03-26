from fastapi import FastAPI, HTTPException

from llama_index.core import (
    VectorStoreIndex, Settings, StorageContext, Document
)
from llama_index.core.schema import TextNode
#from llama_index.llms.openai import OpenAI Removed OpenAI call
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.vector_stores.faiss import FaissVectorStore
from llama_index.vector_stores.postgres import PGVectorStore
from llama_index.core.vector_stores.types import VectorStoreQueryMode
from llama_index.core.indices.keyword_table import KeywordTableIndex
from llama_index.core.indices.composability import ComposableGraph
from llama_index.core.postprocessor import SimilarityPostprocessor
from llama_index.core.node_parser import SentenceSplitter
from pgvector_service import PGVectorService  # Import our Postgres service
from pydantic import BaseModel
import uvicorn
import faiss
from openai import OpenAI
import os
import requests
from github import Github  # PyGithub for private repos
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document as DocxDocument
import pandas as pd
from tika import parser
import io
import uuid
import traceback
import anthropic


# ✅ GitHub Config
GITHUB_TOKEN = ""  # 🔹 Required for private repos
#GITHUB_REPO = "mirkopetracca/genai-automation-elastic-poc"
GITHUB_BRANCH = "main"
#RAW_GITHUB_URL = f"https://raw.githubusercontent.com/{GITHUB_REPO}/{GITHUB_BRANCH}/"

# ✅ OpenAI Config
OPENAI_API_KEY = ""

# ✅ Anthropic Config
ANTHROPIC_API_KEY = ""

# ✅ PGVector Configuration
DB_NAME = "ai_db"
DB_USER = "ai_user"
DB_PASSWORD = "ai_password"
DB_HOST = "localhost"
DB_PORT = "5432"
DB_TABLE_NAME = "document_embeddings"

# ✅ Function to extract text from different file formats
def extract_text_from_file(file_url, file_extension):
    response = requests.get(file_url)
    response.raise_for_status()
    content = response.content

    if file_extension in ["txt", "md"]:
        return content.decode("utf-8", errors="ignore")  # ✅ Plain text

    elif file_extension in ["doc", "docx"]:
        doc = DocxDocument(io.BytesIO(content))  # ✅ Read .docx
        return "\n".join([para.text for para in doc.paragraphs])

    elif file_extension in ["xls", "xlsx", "csv"]:
        df = pd.read_excel(io.BytesIO(content)) if "xls" in file_extension else pd.read_csv(io.BytesIO(content))  # ✅ Read Excel/CSV
        return df.to_string()

    elif file_extension == "pdf":
        with open("temp.pdf", "wb") as temp_file:
            temp_file.write(content)  # ✅ Save temporarily
        text = extract_pdf_text("temp.pdf")
        os.remove("temp.pdf")
        return text

    elif file_extension in ["ppt", "pptx", "uml"]:
        parsed = parser.from_buffer(content)  # ✅ Read PPT/UML
        return parsed["content"] if parsed else ""

    else:
        return None  # ❌ Unsupported file

# ✅ Load Documents from GitHub
#def load_documents_from_github():
#    return load_documents_from_github_repo(GITHUB_REPO)
    
# ✅ Load Documents from GitHub
def load_documents_from_github_repo(github_repo, folder):
    documents = []
    g = Github(GITHUB_TOKEN)
    githubUrl = f"https://raw.githubusercontent.com/{github_repo}/{GITHUB_BRANCH}/"
    repo = g.get_repo(github_repo)

    # 🔹 Get a list of all files in the repo
    contents = repo.get_contents(folder)
    
    # ✅ Postgres Vector Store Connection
    pg_vector_service = PGVectorService(DB_NAME, DB_USER, DB_PASSWORD, DB_HOST, DB_PORT)
    
    existing_embeddings = pg_vector_service.get_existing_embeddings()
    
    
    while contents:
        file_content = contents.pop(0)

        if file_content.type == "dir":
            contents.extend(repo.get_contents(file_content.path))  # ✅ Recursively get subfolders
        else:
            file_url = githubUrl + file_content.path
            file_extension = file_content.path.split(".")[-1].lower()

            text = extract_text_from_file(file_url, file_extension)
            if text and (file_content.name not in existing_embeddings):
                documents.append(Document(text=text, metadata={"file_name": file_content.name, "summary": generate_summary(text)}))

    pg_vector_service.close()
    
    return load_or_create_documents_from_db(documents)

def generate_summary(text: str) -> str:
    #"You are an assistant helping summarize technical documents. "
    #"Summarize the following content with focus on functionality, components, or logic. "
    #"Use simple, clear language.\n\n"
    prompt = (
        "Sei un assistente che aiuta a riassumere documenti tecnici. "
        "Riassumi il contenuto seguente concentrandoti su funzionalità, componenti o logica. "
        "Usa un linguaggio semplice e chiaro.\n\n"
        f"{text[:3000]}"  # truncate to prevent over-tokenizing
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4",  # You can change this to gpt-3.5-turbo if needed
            messages=[
                {"role": "system", "content": "Riassumi i documenti tecnici."}, #You summarize technical documents."
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=512
        )
        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error generating summary: {e}")
        return text[:3000]  # fallback: first part of text

# ✅ Load or create Documents from Postgres and create vector index
def load_or_create_documents_from_db(documents):
    # ✅ Postgres Vector Store Connection
    pg_vector_service = PGVectorService(DB_NAME, DB_USER, DB_PASSWORD, DB_HOST, DB_PORT)
       
    global vector_index
    
    # ✅ Check Existing Embeddings
    existing_embeddings = pg_vector_service.get_existing_embeddings()

    splitter = SentenceSplitter(chunk_size=512, chunk_overlap=50)

    # ✅ Separate New Documents (Batch Processing)
    new_file_names = []
    new_texts = []
    document_embeddings = {}
    summary_map = []
    new_documents = []
    
    for doc in documents:
        file_name = doc.metadata["file_name"]

        # ✅ If embedding exists in DB, use it
        if file_name in existing_embeddings:
            document_embeddings[file_name] = existing_embeddings[file_name]
        else:
            new_documents.append(doc)
            summary = generate_summary(doc.text)
            summary_chunks = splitter.split_text(summary)

            for idx, chunk in enumerate(summary_chunks):
                new_file_names.append(f"{file_name}__chunk_{idx}")
                new_texts.append(chunk)
                summary_map.append(summary)  # or chunk if you want per-chunk summary

    # ✅ Batch Process Only New Documents
    if new_texts:
        print(f"🔹 Generating embeddings for {len(new_texts)} new documents...")
        
        # ✅ Call OpenAI ONCE with all texts
        new_embeddings = Settings.embed_model._get_text_embeddings(new_texts)

        # ✅ Store all embeddings in Postgres in a single transaction
        pg_vector_service.store_embeddings_batch(new_file_names, new_embeddings, summary_map)
              
        vector_index = VectorStoreIndex.from_documents(new_documents, storage_context=storage_context)
        # Optionally, print or log for debugging purposes
        print(f"Added {len(new_file_names)} embeddings to the vector store.")
    
    #node = TextNode(
    #    text=new_texts[0],
    #    metadata={"file_name": new_file_names[0], "summary": summary_map[0]},
    #    node_id=str(uuid.uuid4),
    #    embedding=new_embeddings[0]
    #)
    
    # Add new chunk embedding to the in-memory vector store (VectorStoreIndex)
    #pgvector_store.add(node)

    vector_index = VectorStoreIndex.from_vector_store(storage_context.vector_store)
    #vector_index = VectorStoreIndex.from_documents(documents, storage_context=storage_context)
    
    # ✅ Close DB Connection
    pg_vector_service.close()

    return new_file_names

def generate_code_with_claude(prompt: str, context_chunks: list[str]) -> str:
    client = anthropic.Anthropic()

    context = "\n\n---\n\n".join(context_chunks)
    full_prompt = (
        "You are an expert software engineer. Based on the following technical documentation:\n\n"
        f"{context}\n\n"
        "Write code in response to this task:\n"
        f"{prompt}"
    )

    try:
        response = client.messages.create(
            model="claude-3-opus-20240229",  # or claude-3-sonnet
            max_tokens=1024,
            temperature=0.3,
            messages=[{"role": "user", "content": full_prompt}]
        )
        return response.content[0].text.strip()
    except Exception as e:
        return f"[Error calling ClaudeAI: {str(e)}]"

def get_vector_store():
    # Fetch all existing embeddings from the database.
    global DB_NAME
    vector_store = PGVectorStore.from_params(
        database=DB_NAME, user=DB_USER, password=DB_PASSWORD, host=DB_HOST, port=DB_PORT, table_name=DB_TABLE_NAME
    )
    return vector_store

# ✅ FastAPI API Setup
app = FastAPI()
    
client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    api_key=OPENAI_API_KEY,
)

os.environ["ANTHROPIC_API_KEY"] = ANTHROPIC_API_KEY

# ✅ Use ONLY OpenAI Embeddings (No LLM Calls)
Settings.llm = None #No LLM #llm = OpenAI(model="gpt-4", temperature=0.0, api_key=OPENAI_API_KEY) 
Settings.embed_model = OpenAIEmbedding(model="text-embedding-3-small", api_key=OPENAI_API_KEY)

# ✅ Initialize Postgres Vector Store
pgvector_store = get_vector_store()

# ✅ Index with Existing or Newly Generated Embeddings
storage_context = StorageContext.from_defaults(vector_store=pgvector_store)
    
# ✅ Index with Existing or Newly Generated Embeddings
vector_index = VectorStoreIndex.from_vector_store(storage_context.vector_store)


class QueryRequest(BaseModel):
    query: str

@app.post("/query")
async def query_llama(request: QueryRequest):
    try:
        # ✅ Create Retriever
        retriever = vector_index.as_retriever(similarity_top_k=10)
        response = retriever.retrieve(request.query)
        
        print(response)

        # ✅ Fix: Iterate over response directly (No `.source_nodes`)
        results = [
            {
                "id": i,
                "filename": r.metadata.get("file_name", "Unknown").rpartition("\\")[2],
                "summary": r.metadata.get("summary", ""),  # 🔹 return summary too
                "text": r.text,
                "score": r.score,
            }
            for i, r in enumerate(response)
        ]
        return {"query": request.query, "results": results}

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


class DocumentsRequest(BaseModel):
    githubRepo: str
    folder: str
    
@app.post("/documents")
async def documents_llama(request: DocumentsRequest):
    try:
        response = load_documents_from_github_repo(request.githubRepo, request.folder)
        message = "{} file/s added".format(len(response))
        results = [
            {"id": i, "filename": r}
            for i, r in enumerate(response)
        ]
        return {"githubRepo": request.githubRepo, "message": message, "results": results}

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/clear")
async def clear_llama():
    try:
        vector_store.clear()
        vector_index = VectorStoreIndex.from_vector_store(storage_context.vector_store)
        return {"OK"}

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


class GenerateRequest(BaseModel):
    query: str
    task: str

@app.post("/generate")
async def generate_from_summary(request: GenerateRequest):
    try:
        # Retrieve context summary chunks
        retriever = vector_index.as_retriever(similarity_top_k=10)
        response = retriever.retrieve(request.query)
        top_chunks = [r.text for r in response if r.score and r.score > 0][:5]

        # Send to Claude
        generated_code = generate_code_with_claude(request.task, top_chunks)

        return {
            "task": request.task,
            "code": generated_code,
            "source_chunks": top_chunks
        }

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8001)